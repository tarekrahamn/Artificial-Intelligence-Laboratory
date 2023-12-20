import numpy as np
from docx import Document
import random

# Load the diabetes dataset
diabetes_data = np.genfromtxt(r'C:\\Users\\ASUS\\Desktop\\AiAss\\diabetes.csv', delimiter=',')

def split_dataset(diabetes_data):
    data = []
    for line in diabetes_data:
        exm = line[:-1]
        test = line[-1]
        data.append((exm, test))
    return data

random.shuffle(diabetes_data)
split_data = split_dataset(diabetes_data)

Train, Val, Test = [], [], []

for sample in split_data:
    R = random.uniform(0, 1)
    if R >= 0 and R <= 0.7:
        Train.append(sample)
    elif R > 0.7 and R <= 0.85:
        Val.append(sample)
    else:
        Test.append(sample)


# Split the data into Train, Validation, and Test sets
# Let's assume you have defined the split_data into Train, Val, and Test correctly
# Validation Mean Squared Error Calculation
def knn_regression(Val, Train, K):
    error = 0
    for V in Val:
        distances = []
        for T in Train:
            distance = np.sqrt(np.sum((V[0] - T[0]) ** 2))
            distances.append((T, distance))
        distances.sort(key=lambda x: x[1])
        nearest_neighbors = distances[:K]

        predicted_output = np.mean([sample[0][1] for sample in nearest_neighbors])
        error += (V[1] - predicted_output) ** 2

    mean_squared_error = error / len(Val)
    return mean_squared_error

# Test case to find the best K
def find_best_k(Val, Train):
    K = 5
    min_mse = float('inf')
    for k_value in range(1, 200, 2):
        mse = knn_regression(Val, Train, k_value)
        if mse < min_mse:
            min_mse = mse
            K = k_value
    return K

# Get the best K value
best_K = find_best_k(Val, Train)
print(f"The best K value is {best_K}")

# Create a Word document and add a table
doc = Document()
doc.add_heading('KNN Regression Results', 0)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add column headers to the table
table.cell(0, 0).text = 'K'
table.cell(0, 1).text = 'Mean Squared Error'

# Calculate and print Mean Squared Error for different K values
row_num = 1
for k_value in range(1, 200, 2):
    mse = knn_regression(Val, Train, k_value)
    row_num += 1

    # Extend the table to add enough rows
    while row_num >= len(table.rows):
        table.add_row()

    # Add the K and MSE values to the table
    table.cell(row_num, 0).text = str(k_value)
    table.cell(row_num, 1).text = f'{mse:.10f}'

# Save the document with the results
doc.save('knn_regression_results.docx')

# Calculate and print Mean Squared Error using the best K on the test set
best_mse = knn_regression(Test, Train, best_K)
print(f"Mean Squared Error with best K ({best_K}): {best_mse:.10f}")
