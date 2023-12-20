import numpy as np
import random
from docx import Document

# Load the dataset
dataset = np.genfromtxt(r'C:\Users\ASUS\Desktop\AiAss\iris.csv', delimiter=',')

def split_dataset(dataset):
    data = []
    for line in dataset:
        exm = line[:-1]
        test = line[-1]
        data.append((exm, test))
    return data

random.shuffle(dataset)
split_data = split_dataset(dataset)

Train, Val, Test = [], [], []

for sample in split_data:
    R = random.uniform(0, 1)
    if R >= 0 and R <= 0.7:
        Train.append(sample)
    elif R > 0.7 and R <= 0.85:
        Val.append(sample)
    else:
        Test.append(sample)

# Validation Accuracy Calculation
def knn(Val, Train, K):
    validation_correct = 0
    for V in Val:
        distances = []
        class_votes = {}
        for T in Train:
            distance = np.sqrt(np.sum((np.array(V[0], dtype=float) - np.array(T[0], dtype=float))**2))
            distances.append((T, distance))

        distances.sort(key=lambda x: x[1])
        nearest_neighbors = distances[:K]

        for sample in nearest_neighbors:
            class_L = sample[0][1]
            if class_L in class_votes:
                class_votes[class_L] += 1
            else:
                class_votes[class_L] = 1

        detected_class = max(class_votes, key=class_votes.get)

        if detected_class == V[1]:
            validation_correct += 1

    val_accuracy = (validation_correct / len(Val)) * 100

    return val_accuracy

# Testcase to find the best K
def testcase(Val, Train):
    perfect = 0
    best_K = 0
    for i in range(5, len(Train), 2):
        acc = knn(Val, Train, i)
        if acc > perfect:
            perfect = acc
            best_K = i
    return best_K

# Get the best K value
K = testcase(Val, Train)
print(f"The best K value is {K}")

# Create a Word document and add a table
doc = Document()
doc.add_heading('K-NN Results', 0)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add column headers to the table
table.cell(0, 0).text = 'K'
table.cell(0, 1).text = 'Val_acc'
# Calculate and print Validation Accuracy for different K values
row_num = 1
for k_value in range(1, 15, 2):
    val_accuracy = knn(Val, Train, k_value)
    test_accuracy = knn(Test,Train, k_value)
    doc.add_paragraph(f' Test Accuracy with Best K={k_value} : {val_accuracy:.2f}%')
    row_num += 1


    # Extend the table to add enough rows
    while row_num >= len(table.rows):
        table.add_row()

    # Add the K and Val_acc values to the table
    table.cell(row_num, 0).text = str(k_value)
    table.cell(row_num, 1).text = f'{val_accuracy:.2f}%'

# Save the document
doc.save('knn_results.docx')

# Calculate and print Test Accuracy using the best K
test_accuracy = knn(Test, Train, K)
print(f"Test Accuracy with best K: {test_accuracy:.2f}%")
